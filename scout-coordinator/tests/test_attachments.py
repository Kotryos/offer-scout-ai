from io import BytesIO

from docx import Document
import pytest
from reportlab.pdfgen import canvas

from scout_coordinator.attachments.extractor import extract_attachment_text, html_to_text


def test_extracts_docx_text() -> None:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Oferta testowa")
    document.add_paragraph("Java, praca zdalna")
    document.save(buffer)

    result = extract_attachment_text(
        filename="offer.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=buffer.getvalue(),
    )

    assert not result.skipped
    assert "Oferta testowa" in result.text
    assert "Java, praca zdalna" in result.text


def test_extracts_pdf_text() -> None:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(100, 750, "Oferta PDF Java")
    pdf.save()

    result = extract_attachment_text(
        filename="offer.pdf",
        content_type="application/pdf",
        data=buffer.getvalue(),
    )

    assert not result.skipped
    assert "Oferta PDF Java" in result.text


def test_converts_html_to_text() -> None:
    assert html_to_text("<h1>Oferta</h1><p>Praca zdalna</p>") == "Oferta\nPraca zdalna"


def test_extracts_plain_text_with_charset_content_type() -> None:
    result = extract_attachment_text(
        filename="offer.txt",
        content_type="text/plain; charset=utf-8",
        data="wynagrodzenie, B2B, praca zdalna, ą ć ę ł ń ó ś ź ż".encode(),
    )

    assert not result.skipped
    assert "wynagrodzenie, B2B, praca zdalna" in result.text
    assert "ą ć ę ł ń ó ś ź ż" in result.text


def test_extracts_markdown_text() -> None:
    result = extract_attachment_text(
        filename="offer.md",
        content_type="text/markdown",
        data="# Oferta\n- Java\n- B2B".encode(),
    )

    assert not result.skipped
    assert "# Oferta" in result.text
    assert "- Java" in result.text


def test_extracts_html_attachment_text() -> None:
    result = extract_attachment_text(
        filename="offer.html",
        content_type="text/html",
        data="<h1>Oferta</h1><p>Praca zdalna</p>".encode(),
    )

    assert not result.skipped
    assert result.text == "Oferta\nPraca zdalna"


def test_handles_uppercase_content_type() -> None:
    result = extract_attachment_text(
        filename="offer.txt",
        content_type="TEXT/PLAIN",
        data=b"Java",
    )

    assert not result.skipped
    assert result.text == "Java"


def test_marks_unsupported_attachment_as_skipped() -> None:
    result = extract_attachment_text(
        filename="image.png",
        content_type="image/png",
        data=b"not relevant",
    )

    assert result.skipped
    assert result.reason == "unsupported content type"


@pytest.mark.parametrize(
    ("filename", "content_type"),
    [
        ("broken.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        ("broken.pdf", "application/pdf"),
    ],
)
def test_marks_corrupted_supported_attachment_as_skipped(filename: str, content_type: str) -> None:
    result = extract_attachment_text(
        filename=filename,
        content_type=content_type,
        data=b"not a valid document",
    )

    assert result.skipped
    assert result.reason == "could not extract text"
